import docx

#docFile=open("multipleParagraphs.docs",'rb')

docFile=docx.Document("multipleParagraphs.docx")

print(len(docFile.paragraphs))

fulltext=[]
for para in docFile.paragraphs:
    fulltext.append(para.text)

print('\n'.join(fulltext))


doc1_file=docx.Document()
doc1_file.add_paragraph('/n'.join(fulltext))

doc1_file.save("copyofpara.docx")

print(len(doc1_file.paragraphs))
